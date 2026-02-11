import os
import glob
import faiss
import numpy as np
import hashlib
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import multiprocessing

class DataLoader:
    def __init__(self, data_dir, vector_db_path):
        self.data_dir = data_dir
        self.vector_db_path = vector_db_path
        self.vector_store = None
        self.documents = []
        self.file_hashes = {}
        self.vector_dim = 128  # 固定向量维度
    
    def split_single_text(self, text, chunk_size=500, chunk_overlap=50):
        """简单的文本切分功能（优化版本）"""
        print(f"开始切分文本，长度: {len(text)} 字符")
        chunks = []
        start = 0
        text_length = len(text)
        max_iterations = (text_length + chunk_size - 1) // (chunk_size - chunk_overlap) + 10
        iteration_count = 0
        
        while start < text_length and iteration_count < max_iterations:
            iteration_count += 1
            end = min(start + chunk_size, text_length)
            
            # 简化的句子边界分割
            if end < text_length:
                # 只查找最近的几个标点符号，减少计算量
                punctuation = ['.', '。', '!', '！', '?', '？', ';', '；', '\n']
                # 从end开始向前搜索最多20个字符
                search_end = max(start, end - 20)
                for i in range(end - 1, search_end - 1, -1):
                    if text[i] in punctuation:
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 计算下一个起始位置
            next_start = end - chunk_overlap
            # 防止无限循环
            if next_start <= start:
                next_start = start + 1
            start = next_start
        
        if iteration_count >= max_iterations:
            print(f"警告：文本切分达到最大迭代次数 {max_iterations}，可能存在问题")
        
        print(f"文本切分完成，生成 {len(chunks)} 个 chunks")
        return chunks
    
    def simple_embed(self, text):
        """简单的基于字符频率的文本嵌入方法"""
        # 统计常见字符的频率
        char_freq = {}
        for char in text:
            if char in char_freq:
                char_freq[char] += 1
            else:
                char_freq[char] = 1
        
        # 生成固定长度的向量
        vector = np.zeros(self.vector_dim, dtype=np.float32)
        
        # 使用字符的ASCII码作为索引，将频率映射到向量
        for i, (char, freq) in enumerate(char_freq.items()):
            if i < self.vector_dim:
                vector[i] = freq / len(text)  # 归一化频率
        
        return vector
    
    def load_files(self):
        """加载指定目录下的所有文件"""
        files = glob.glob(os.path.join(self.data_dir, "**/*"), recursive=True)
        valid_files = []
        
        for file_path in files:
            if os.path.isfile(file_path):
                try:
                    content = self._read_file(file_path)
                    if content:
                        valid_files.append((file_path, content))
                except Exception as e:
                    print(f"无法读取文件 {file_path}: {str(e)}")
        
        return valid_files
    
    def _read_file(self, file_path):
        """根据文件类型读取文件内容"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt' or ext == '.md' or ext == '.csv':
            # 读取文本文件
            print(f"正在读取文本文件: {file_path}")
            # 使用生成器方式分批读取
            content_generator = self._read_large_file(file_path, chunk_size=1024*1024)  # 1MB chunks
            content = ''.join(content_generator)
            print(f"读取完成，文本长度: {len(content)} 字符")
            # 限制文本长度，避免处理过大的文件
            max_length = 500000  # 500KB
            if len(content) > max_length:
                print(f"文本过长，截断到 {max_length} 字符")
                content = content[:max_length]
            return content
        elif ext == '.docx':
            # 读取Word文档
            print(f"正在读取Word文档: {file_path}")
            # 限制Word文档处理规模
            content = self._read_large_docx(file_path, max_paras=500, max_length=500000)
            if content:
                print(f"读取完成，文本长度: {len(content)} 字符")
            return content
        else:
            # 不支持的文件类型
            print(f"不支持的文件类型: {ext}")
            return None
    
    def _read_large_file(self, file_path, chunk_size=1024*1024):
        """分批读取大文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                yield chunk
    
    def _read_large_docx(self, file_path, max_paras=500, max_length=500000):
        """分批处理大Word文档"""
        try:
            doc = Document(file_path)
            full_text = []
            para_count = len(doc.paragraphs)
            print(f"文档包含 {para_count} 个段落")
            
            # 限制处理的段落数量
            actual_paras = min(para_count, max_paras)
            if para_count > max_paras:
                print(f"段落过多，只处理前 {actual_paras} 个段落")
            
            current_length = 0
            for i in range(actual_paras):
                if i % 100 == 0:
                    print(f"处理段落 {i}/{actual_paras}")
                
                para_text = doc.paragraphs[i].text
                para_length = len(para_text)
                
                # 检查是否超过长度限制
                if current_length + para_length > max_length:
                    print(f"文本长度达到限制，停止处理")
                    break
                
                full_text.append(para_text)
                current_length += para_length
            
            content = '\n'.join(full_text)
            return content
        except Exception as e:
            print(f"处理Word文档时出错: {str(e)}")
            return None
    
    def split_text(self, files):
        """切分文本为小块"""
        documents = []
        print(f"开始处理 {len(files)} 个文件的文本切分")
        
        for file_path, content in files:
            print(f"处理文件: {file_path}")
            chunks = self.split_single_text(content)
            for i, chunk in enumerate(chunks):
                doc = {
                    "page_content": chunk,
                    "metadata": {
                        "source": file_path,
                        "chunk_index": i
                    }
                }
                documents.append(doc)
            print(f"文件处理完成: {file_path}, 生成 {len(chunks)} 个文档块")
        
        print(f"文本切分完成，共生成 {len(documents)} 个文档块")
        return documents
    
    def compute_embeddings(self, documents):
        """计算文本嵌入向量"""
        print(f"开始计算 {len(documents)} 个文档的嵌入向量")
        
        embeddings = []
        batch_size = 50  # 批处理大小
        total_batches = (len(documents) + batch_size - 1) // batch_size
        
        # 分批处理
        for batch_idx in range(total_batches):
            start_idx = batch_idx * batch_size
            end_idx = min((batch_idx + 1) * batch_size, len(documents))
            batch_docs = documents[start_idx:end_idx]
            
            print(f"处理批次 {batch_idx + 1}/{total_batches}, 文档范围: {start_idx}-{end_idx}")
            
            # 处理当前批次
            for doc in batch_docs:
                text = doc["page_content"]
                embedding = self.simple_embed(text)
                embeddings.append(embedding)
        
        print(f"嵌入向量计算完成，共生成 {len(embeddings)} 个向量")
        return np.array(embeddings).astype('float32')
    
    def build_vector_store(self, documents, embeddings):
        """构建FAISS向量库（优化版本）"""
        print(f"开始构建向量库，文档数: {len(documents)}, 向量维度: {embeddings.shape[1]}")
        
        dimension = embeddings.shape[1]
        
        # 根据向量数量选择合适的索引类型
        if len(embeddings) < 1000:
            # 小数据集使用简单索引
            print("使用 IndexFlatL2 索引（小数据集）")
            index = faiss.IndexFlatL2(dimension)
        else:
            # 大数据集使用IVF索引加速
            nlist = min(100, len(embeddings) // 10)
            print(f"使用 IndexIVFFlat 索引（大数据集），聚类数: {nlist}")
            quantizer = faiss.IndexFlatL2(dimension)
            index = faiss.IndexIVFFlat(quantizer, dimension, nlist, faiss.METRIC_L2)
            # IVF索引需要训练
            print("训练IVF索引...")
            index.train(embeddings)
        
        # 添加向量
        print("添加向量到索引...")
        index.add(embeddings)
        print(f"向量库构建完成，共添加 {index.ntotal} 个向量")
        
        # 保存向量库和文档
        print(f"保存向量库到: {self.vector_db_path}")
        faiss.write_index(index, self.vector_db_path)
        self.vector_store = index
        self.documents = documents
        
        # 更新文件哈希
        self._update_file_hashes()
        print("向量库构建和保存完成")
        
    def load_vector_store(self):
        """加载已有的向量库"""
        if os.path.exists(self.vector_db_path):
            try:
                self.vector_store = faiss.read_index(self.vector_db_path)
                # 重新加载文档（实际应用中可能需要更高效的存储方式）
                files = self.load_files()
                self.documents = self.split_text(files)
                self._update_file_hashes()
                return True
            except Exception as e:
                print(f"加载向量库失败: {str(e)}")
                return False
        return False
    
    def _update_file_hashes(self):
        """更新文件哈希值，用于检测文件变更"""
        files = glob.glob(os.path.join(self.data_dir, "**/*"), recursive=True)
        for file_path in files:
            if os.path.isfile(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        file_hash = hashlib.md5(f.read()).hexdigest()
                    self.file_hashes[file_path] = file_hash
                except Exception as e:
                    print(f"计算文件哈希失败 {file_path}: {str(e)}")
    
    def check_for_changes(self):
        """检查是否有文件变更"""
        current_hashes = {}
        files = glob.glob(os.path.join(self.data_dir, "**/*"), recursive=True)
        
        # 检查新增或修改的文件
        for file_path in files:
            if os.path.isfile(file_path):
                try:
                    with open(file_path, 'rb') as f:
                        file_hash = hashlib.md5(f.read()).hexdigest()
                    current_hashes[file_path] = file_hash
                    
                    if file_path not in self.file_hashes or self.file_hashes[file_path] != file_hash:
                        return True
                except Exception as e:
                    print(f"检查文件变更失败 {file_path}: {str(e)}")
        
        # 检查删除的文件
        for file_path in self.file_hashes:
            if not os.path.exists(file_path):
                return True
        
        return False
    
    def rebuild_vector_store(self):
        """重新构建向量库"""
        print("检测到文件变更，正在重新构建向量库...")
        files = self.load_files()
        documents = self.split_text(files)
        if documents:
            embeddings = self.compute_embeddings(documents)
            self.build_vector_store(documents, embeddings)
            print("向量库重建完成")
        else:
            print("没有可处理的文档")
