from docx import Document
import os

def extract_doc_content(file_path):
    """提取Word文档内容"""
    try:
        doc = Document(file_path)
        full_text = []
        
        print(f"文档包含 {len(doc.paragraphs)} 个段落")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if text.strip():
                full_text.append(f"段落 {i+1}: {text}")
                # 检查是否包含"出差申请"相关内容
                if "出差申请" in text:
                    print(f"找到相关内容 - 段落 {i+1}: {text}")
        
        content = '\n'.join(full_text)
        return content
    except Exception as e:
        print(f"处理Word文档时出错: {str(e)}")
        return None

if __name__ == "__main__":
    doc_path = "d:\\AI coding- Talk Agent\\审核后资料文件夹\\小天元手册-2026年1月版-费控商旅.docx"
    if os.path.exists(doc_path):
        content = extract_doc_content(doc_path)
        if content:
            # 保存提取的内容到文本文件
            output_path = "d:\\AI coding- Talk Agent\\doc_content.txt"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"文档内容已提取到: {output_path}")
    else:
        print(f"文档不存在: {doc_path}")
